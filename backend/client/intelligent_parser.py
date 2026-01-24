"""Intelligent contract parser with reasoning and web search capabilities.

This parser uses a smart agent (Claude/GPT-4) that:
- Reads and understands the entire contract
- Reasons about staffing patterns and evolution over time
- Uses web search only when document lacks data
- Extracts year-specific staffing intelligently

Supported file types: PDF, DOCX, XLSX, XLS, TXT, CSV
"""

import json
from typing import Dict, Any, List, Optional
from pathlib import Path

from agno.agent import Agent
from agno.tools.reasoning import ReasoningTools
from agno.tools.exa import ExaTools
from app.settings import settings
from client.llm_client import get_chat_llm_agno


def create_reasoning_tool(
    instructions: str = "only show reasoning no need for for action confidence",
    add_instructions: bool = True,
    think: bool = True,
    analyze: bool = True,
    add_few_shot: bool = False,
    few_shot_examples: Optional[List[Dict[str, str]]] = None,
) -> ReasoningTools:
    """
    Create a reasoning tool with customizable parameters.

    Args:
        instructions: Instructions for reasoning
        add_instructions: Whether to add instructions
        think: Enable thinking
        analyze: Enable analysis
        add_few_shot: Whether to add few-shot examples
        few_shot_examples: List of few-shot examples

    Returns:
        Configured ReasoningTools instance
    """
    try:
        tool = ReasoningTools(
            instructions=instructions,
            add_instructions=add_instructions,
            enable_think=think,
            enable_analyze=analyze,
            add_few_shot=add_few_shot,
            few_shot_examples=few_shot_examples,
        )
    
        return tool
    except Exception as e:
        raise RuntimeError(f"Failed to create reasoning tool: {str(e)}")


def _extract_text_from_pdf(file_path: str) -> str:
    """Extract full text from PDF."""
    try:
        import PyPDF2
        pages = []
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text = page.extract_text()
                if text.strip():
                    pages.append(text)
        return '\n\n--- PAGE BREAK ---\n\n'.join(pages)
    except ImportError:
        raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")


def _extract_text_from_docx(file_path: str) -> str:
    """Extract full text from DOCX."""
    try:
        import docx
        doc = docx.Document(file_path)
        content = []

        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)

        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    table_text.append(row_text)
            if table_text:
                content.append('\n'.join(table_text))

        return '\n\n'.join(content)
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")


def _extract_text_from_excel(file_path: str) -> str:
    """Extract full text from Excel (.xlsx, .xls)."""
    try:
        import pandas as pd
        import openpyxl  # noqa: F401 - Required by pandas for Excel support

        # Read all sheets
        excel_file = pd.ExcelFile(file_path)
        content = []

        for sheet_name in excel_file.sheet_names:
            content.append(f"=== SHEET: {sheet_name} ===")

            # Read sheet
            df = pd.read_excel(file_path, sheet_name=sheet_name)

            # Convert to text format (preserve structure)
            # Replace NaN with empty string
            df = df.fillna('')

            # Convert to string representation (table-like format)
            sheet_text = df.to_string(index=False)
            content.append(sheet_text)
            content.append("")  # Add blank line between sheets

        return '\n\n'.join(content)
    except ImportError as e:
        if 'openpyxl' in str(e):
            raise ImportError("openpyxl not installed. Run: pip install openpyxl")
        else:
            raise ImportError(f"Required library not installed: {e}")


def _extract_text(file_path: str) -> str:
    """Extract text from file based on extension."""
    import os
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.pdf':
        return _extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return _extract_text_from_docx(file_path)
    elif ext in ['.xlsx', '.xls']:
        return _extract_text_from_excel(file_path)
    elif ext in ['.txt', '.csv']:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: PDF, DOCX, XLSX, XLS, TXT, CSV")


def _create_intelligent_parser() -> Agent:
    """Create intelligent parser with reasoning and web search capabilities."""

    # Use a powerful model that can reason
    llm=get_chat_llm_agno(model="anthropic/claude-sonnet-4.5",api_key=settings.OPENROUTER_API_KEY,base_url="https://openrouter.ai/api/v1",max_tokens=32000,temperature=0.1)


    # Create reasoning tool with few-shot example for combined team pattern
    few_shot_example = [
        {
            "input": "Base Year: 2-3 analysts combining threat intelligence, threat hunting, and analytics. Option Years: dedicated threat intelligence team of 3-5, threat hunting team of 2-4, analytics team of 5-10.",
            "reasoning": "The document describes a COMBINED team in Base Year (one team doing multiple functions), then SPECIALIZED teams in Option Years (separate dedicated teams). I should create either: (1) a 'Combined CTI/Hunting/Analytics Analyst' position for Base Year with 2.5 FTEs, then separate positions appearing in Option Years, OR (2) distribute the 2.5 Base Year FTEs proportionally across the three specialist roles that appear in Option Years."
        }
    ]

    reasoning_tool = create_reasoning_tool(
        add_few_shot=True,
        few_shot_examples=few_shot_example
    )

    # Web search tool for research (only when needed)
    exa_tool = ExaTools(
        api_key=settings.EXA_API_KEY,
        num_results=10,
        text=True,
        text_length_limit=2000,
        livecrawl="always",
    )

    # Minimal, high-level instructions (like Claude Code's system prompt)
    instructions = [
        """You are an expert government contract analyst.

        WORKFLOW:
        1. FIRST: Use reasoning_tool to analyze the document
           - Read and understand the full contract
           - Identify what data is present and what is missing
           - Recognize patterns in how staffing evolves
           - Decide if web search is needed

        2. IF NEEDED: Use web search to research missing information
           - Only when document lacks specific data
           - Example: typical staffing levels for 24/7 operations
           - Prioritize document data over web research

        3. THEN: Extract structured data as JSON

        Core principles:
        - Document data is the primary source
        - Understand context before extracting
        - Recognize when staffing varies by year
        - Use judgment for ranges and ambiguity
        - Infer implicit information

        Extract intelligently."""
    ]

    agent = Agent(
        name="Intelligent Contract Parser",
        model=llm,
        tools=[reasoning_tool, exa_tool],  # Reasoning + web search
        instructions=instructions,
        markdown=False,
        debug_mode=True
    )

    return agent


async def parse_document_intelligent(
    file_path: str,
    default_fte_hours: int = 1920,
    default_years: int = 5
) -> Dict[str, Any]:
    """
    Parse document with intelligent agent that reasons first, then extracts.

    The agent:
    - Reads the entire contract and understands context
    - Recognizes when staffing varies by year
    - Uses web search only when needed
    - Extracts year-specific staffing intelligently

    Args:
        file_path: Path to document (PDF, DOCX, XLSX, XLS, TXT, CSV)
        default_fte_hours: Default FTE hours per year
        default_years: Default contract duration

    Returns:
        Dict with metadata, positions, travel, odcs
    """
    print(f"\n{'='*70}")
    print(f"Intelligent Parser - Context-Aware Extraction")
    print(f"{'='*70}")
    print(f"\nFile: {Path(file_path).name}")

    # Extract text
    print(f"\n📄 Step 1: Extracting document text...")
    text = _extract_text(file_path)
    print(f"  ✓ Extracted {len(text):,} characters")

    # Create agent
    print(f"\n🤖 Step 2: Creating intelligent agent...")
    print(f"  Tools: reasoning_tool, web_search")
    agent = _create_intelligent_parser()

    # Build prompt (simple, high-level)
    prompt = f"""
Read this government contract document and extract the complete staffing plan.

DOCUMENT:
{text}

WORKFLOW:

1. Use reasoning_tool to analyze:
   - What is the contract structure (base years, option years, extensions)?
   - How is staffing described (tables, narratives, or both)?
   - Are there positions that scale up over time?
   - Are there roles that appear only in later years?
   - Are there "combined teams" that later split into specialized roles?
   - Is there an overtime (OT) provision or multiplier mentioned?
   - Is there a surge option or capacity increase clause?
   - Are surge positions listed separately or is it just a percentage?
   - How can you identify which positions are surge vs base positions?
   - What data is present vs missing?
   - Do you need web search to fill gaps?

2. If needed, use web search to research:
   - Typical staffing levels for operations described
   - Industry standards for specific roles
   - Only when document lacks specific data

3. Extract as JSON with this structure:
{{
  "metadata": {{
    "project_name": "string or null",
    "location": "State name or null",
    "base_years": 1,
    "option_years": 4,
    "total_years": 5,
    "standard_fte_hours": {default_fte_hours},
    "notes": "Add any relevant contract notes "
  }},

  "positions": [
    {{
      "labor_category": "Job title",
      "description": "Full job description",
      "experience": 5,
      "location": "California",  # REQUIRED: Actual state name from document (e.g., "California", "Virginia", "Texas")
      "location_type": "On-Site", # REQUIRED: MUST be exactly "On-Site" or "Off-Site" (no "Remote", "Hybrid", etc.)
      "is_key_position": false,
      "is_surge": false,  # REQUIRED: true if this is a surge position, false if base position
      "hours_per_year": {{
        "1": 1920,    # REQUIRED: Year-specific regular hours. Use "1", "2", "3", etc. as keys
        "2": 1920,    # If hours are constant: repeat same value for all years
        "3": 1920,
        "4": 1920,
        "5": 1920,
        "6": 960      # CRITICAL: If extensions exist (year 6 in this case), MUST include extension year hours! Prorate based on duration_months if not specified.
      }},
      "ot_hours_per_year": {{
        "1": 200,     # OPTIONAL: Overtime hours per year (if mentioned in document)
        "2": 200,     # Set to 0 or omit years with no OT
        "3": 0,
        "4": 0,
        "5": 0
      }},
      "data_source": "document" or "web_research"
    }}
  ],

  "travel": [
    {{
      "description": "Travel description",
      "amount_per_year": {{"1": 50000, "2": 51500}}
    }}
  ],

  "odcs": [
    {{
      "category": "Equipment",
      "description": "Description",
      "amount_per_year": {{"1": 10000, "2": 10300}}
    }}
  ],

  "extensions": [
    {{
      "year": 6,  # Extension year number (beyond total_years)
      "label": "6 Month Extension",  # Display label (e.g., "6 Month Extension", "12 Month Extension")
      "duration_months": 6,  # Duration in months (typically 6 or 12)
      "description": "Extension Period 6"  # Description or notes about this extension
    }}
  ],

  "surge": {{
    "percentage": 0.20,  # Decimal form (20% = 0.20). Set to null if no surge option.
    "description": "Government has option to increase contract by up to 20%"  # Full context from document
  }}
}}

IMPORTANT:
- Prioritize document data over web research
- Use web search only for missing information
- Document your reasoning and data sources
- For FTE ranges (e.g., "2-3 analysts"), use midpoint or explain your choice
- For positions appearing later (e.g., "Option Years only"), set early years to 0
- Watch for "combined teams" (Base Year) that split into specialized roles (Option Years)
- CRITICAL: Always use year-specific format for hours_per_year: {{"1": hours, "2": hours, ...}}
- If hours are constant across all years: repeat the same value for each year key
- Never use "all_years" or similar keys - always use numeric year keys ("1", "2", "3", etc.)
- CRITICAL: location_type MUST be exactly "On-Site" or "Off-Site":
  * Use "On-Site" if: work at government facility, client site, or document says "on-site"
  * Use "Off-Site" if: remote work, contractor facility, telecommute, or document says "remote"/"offsite"
  * Never use "Remote", "Hybrid", "Telecommute", or any other value
- CRITICAL: location MUST be actual state name from document:
  * Extract the actual state where work is performed (e.g., "California", "Virginia", "Maryland")
  * Look for state names in contract location/place of performance sections
  * If document says "remote" or doesn't specify state, use "National"
  * Never use placeholder text like "State name" or "TBD"
- Extensions are optional periods beyond the regular contract (base + option years):
  * Look for phrases like "6-month extension", "12-month extension", "extension period"
  * If document mentions year 6 or beyond when total_years is 5, that's an extension
  * Extract year number, duration in months, and any description
  * Leave empty array [] if no extensions mentioned
- CRITICAL: When extensions exist, positions MUST include extension year hours in hours_per_year:
  * If document specifies extension hours, use those values
  * If not specified, prorate based on duration_months: (duration_months / 12) × base_year_hours
  * Example: 6-month extension → year_6_hours = (6 / 12) × year_5_hours = 0.5 × year_5_hours
  * Example: 12-month extension → year_6_hours = year_5_hours (same as regular year)
  * NEVER leave extension years out of hours_per_year dict!
- Overtime (OT) hours extraction (optional - only extract if mentioned in document):
  * Look for phrases like: "overtime hours", "OT", "beyond 40 hours/week", "additional hours"
  * Extract OT hours per position per year in ot_hours_per_year field (same format as hours_per_year)
  * If document specifies OT hours by position and year, extract them
  * If NO OT hours mentioned, omit ot_hours_per_year field entirely (do not set to 0)
  * ONLY extract OT HOURS - do NOT extract multipliers (like "time-and-a-half", "1.5x", "double-time")
- Surge option extraction (TWO SCENARIOS - extract what's in the document):
  * SCENARIO 1: Surge Positions (specific labor categories with hours)
    - Look for separate staffing sections: "Surge Staffing", "Surge Positions", "Surge Capacity"
    - Look for position names indicating surge: "Surge [Role]", "[Role] (Surge)", "[Role] - Surge"
    - Look for phrases: "surge positions", "additional surge staff", "surge labor"
    - For these positions, set is_surge: true in the position object
    - Common pattern: document lists base staffing, then separate section for surge staffing
    - Example: "Base: 5 Engineers, Surge: 3 additional Engineers" → extract 5 base + 3 surge positions
  * SCENARIO 2: Percentage-based Surge (no specific positions, just %)
    - Look for phrases like: "surge option", "increase contract by X%", "up to X% additional capacity"
    - Look for DFARS clause references: "252.217-7001", "FAR 52.217-7", "surge clause"
    - Extract percentage as DECIMAL: 20% → 0.20, 50% → 0.50
    - Include full description/context from document
    - Common surge percentages: 20%, 25%, 50% (up to 50% is standard, over 50% is rare)
  * CRITICAL RULES:
    - If document has BOTH surge positions AND percentage, extract positions (Scenario 1 takes precedence)
    - If document ONLY has percentage, set surge.percentage and all positions get is_surge: false
    - If document has NEITHER, set surge: null and all positions get is_surge: false
    - ONLY extract surge PERCENTAGE - do NOT extract multipliers, premiums, or pricing details
    - Default all positions to is_surge: false unless clear evidence they are surge positions

Return ONLY valid JSON, no markdown code blocks.
"""

    # Run agent with retry logic for malformed JSON
    print(f"\n🔍 Step 3: Running agent (will reason, search if needed, then extract)...\n")

    max_retries = 3
    retry_count = 0
    result = None

    while retry_count < max_retries:
        try:
            # Add retry context to prompt if this is a retry
            current_prompt = prompt
            if retry_count > 0:
                current_prompt = f"{prompt}\n\nIMPORTANT: Your previous response was truncated or malformed. Please provide COMPLETE valid JSON. Ensure all brackets and quotes are closed properly."
                print(f"  🔄 Retry attempt {retry_count}/{max_retries - 1} due to malformed JSON...")

            response = await agent.arun(current_prompt)

            # Parse JSON response
            response_text = response.content if hasattr(response, 'content') else str(response)

            # Clean up markdown if present
            if '```json' in response_text:
                response_text = response_text.split('```json')[1].split('```')[0]
            elif '```' in response_text:
                response_text = response_text.split('```')[1].split('```')[0]

            response_text = response_text.strip()

            # Try to parse JSON
            result = json.loads(response_text)
            print(f"  ✅ JSON parsed successfully")
            break  # Success! Exit retry loop

        except json.JSONDecodeError as e:
            print(f"\n  ❌ JSON parsing error (attempt {retry_count + 1}/{max_retries}): {e}")
            print(f"  Response preview: {response_text[:500] if 'response_text' in locals() else 'N/A'}...")
            retry_count += 1

            # If we've exhausted retries, return empty result
            if retry_count >= max_retries:
                print(f"  ❌ Failed after {max_retries} attempts. Returning empty result.")
                result = {
                    'metadata': {
                        'total_years': default_years,
                        'standard_fte_hours': default_fte_hours
                    },
                    'positions': [],
                    'travel': [],
                    'odcs': [],
                    'extensions': [],
            'surge': None
        }

    # Summary
    print(f"\n{'='*70}")
    print(f"✅ Extraction Complete")
    print(f"{'='*70}")

    metadata = result.get('metadata', {})
    positions = result.get('positions', [])
    travel = result.get('travel', [])
    odcs = result.get('odcs', [])
    extensions = result.get('extensions', [])
    surge = result.get('surge', None)

    print(f"\n  Project: {metadata.get('project_name', 'Unknown')}")
    print(f"  Location: {metadata.get('location', 'Unknown')}")
    print(f"  Duration: {metadata.get('total_years', default_years)} years")
    if extensions:
        print(f"  Extensions: {len(extensions)} period(s)")

    # Surge detection summary
    surge_positions = [p for p in positions if p.get('is_surge', False)]
    if surge_positions:
        print(f"  Surge: {len(surge_positions)} surge position(s) detected (Scenario 1: Specific positions)")
    elif surge and surge.get('percentage'):
        print(f"  Surge: {surge.get('percentage') * 100:.1f}% option (Scenario 2: Percentage-based) - {surge.get('description', 'No description')}")

    print(f"\n  Positions extracted: {len(positions)} total")
    if surge_positions:
        print(f"    - Base positions: {len(positions) - len(surge_positions)}")
        print(f"    - Surge positions: {len(surge_positions)}")
    print(f"  Travel items: {len(travel)}")
    print(f"  ODC items: {len(odcs)}")

    if positions:
        print(f"\n  Position samples:")
        for i, pos in enumerate(positions[:5]):
            labor_cat = pos.get('labor_category', 'Unknown')
            hours_y1 = pos.get('hours_per_year', {}).get('1', 0)
            source = pos.get('data_source', 'document')
            is_surge = pos.get('is_surge', False)
            surge_tag = " [SURGE]" if is_surge else ""
            print(f"    {i+1}. {labor_cat[:50]}{surge_tag} (Year 1: {hours_y1} hrs, Source: {source})")

        if len(positions) > 5:
            print(f"    ... and {len(positions) - 5} more positions")

    print()
    return result


def test_intelligent_parser(file_path: str) -> Dict[str, Any]:
    """Test the intelligent parser on a file."""
    import asyncio
    return asyncio.run(parse_document_intelligent(file_path))


if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        result = test_intelligent_parser(sys.argv[1])
        print(f"\n\n=== FINAL JSON ===")
        print(json.dumps(result, indent=2, default=str))
    else:
        print("Usage: python intelligent_parser.py <path_to_document>")
