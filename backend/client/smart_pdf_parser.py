"""Smart PDF parsing using Agent with Exa web search for research.

This parser uses an Agno Agent with web search capability to parse documents
and research missing data like FTE counts and staffing levels.
"""

import json
from typing import List, Dict, Any
from pathlib import Path

from agno.agent import Agent
from agno.models.openai import OpenAIChat
from agno.tools.exa import ExaTools
from app.settings import settings


def _extract_text_from_pdf(file_path: str) -> str:
    """Extract full text from PDF."""
    try:
        import PyPDF2
        pages = []
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text = page.extract_text()
                if text.strip():
                    pages.append(text)
        return '\n\n--- PAGE BREAK ---\n\n'.join(pages)
    except ImportError:
        raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")


def _extract_text_from_docx(file_path: str) -> str:
    """Extract full text from DOCX."""
    try:
        import docx
        doc = docx.Document(file_path)
        content = []

        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)

        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    table_text.append(row_text)
            if table_text:
                content.append('\n'.join(table_text))

        return '\n\n'.join(content)
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")


def _extract_text(file_path: str) -> str:
    """Extract text from file based on extension."""
    import os
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.pdf':
        return _extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return _extract_text_from_docx(file_path)
    elif ext in ['.txt', '.csv']:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _create_parsing_agent() -> Agent:
    """Create an agent with Exa web search for document parsing."""

    # Create LLM
    llm = OpenAIChat(
        id="openai/gpt-4.1-mini",
        api_key=settings.OPENROUTER_API_KEY,
        base_url="https://openrouter.ai/api/v1",
        max_tokens=50000,
        temperature=0,
    )

    # Create Exa web search tool
    exa_tool = ExaTools(
        api_key=settings.EXA_API_KEY,
        num_results=10,
        text=True,
        text_length_limit=2000,
        highlights=True,
        livecrawl="always",
    )

    instructions = [
        """You are a Government Contract Staffing Analyst Agent. Your task is to parse documents and extract ALL staffing information.

CRITICAL RULE: ALWAYS extract data from the document FIRST. DO NOT use web search if data exists in the document.

STEP 1 - SEARCH THE DOCUMENT FOR:
- Tables with labor hours by year (e.g., "Year 1", "Year 2", "Base Period", "Option Year")
- FTE counts or headcount numbers
- Level of Effort (LOE) data
- Staffing matrices or labor mix tables
- Any numerical staffing data

STEP 2 - EXTRACTION RULES:
- If you find hours/FTE data in the document → Extract EXACT numbers, DO NOT estimate
- If hours are in a table → Extract each year's hours exactly as shown
- If only total hours given → Distribute across years
- Calculate FTEs from hours: FTEs = hours ÷ 1920 (or document's standard hours)

STEP 3 - WEB SEARCH (ONLY AS LAST RESORT):
Use Exa web search ONLY when ALL of these are true:
1. Document has NO hours table
2. Document has NO FTE counts
3. Document has NO level of effort data
4. Document only describes job qualifications without quantities""",

        """For each position, you MUST provide:
- labor_category: Job title (REQUIRED) - extract exactly as written
- description: Job duties - extract from document
- experience: Years required (integer) - extract from document
- location: State name
- location_type: "On-Site" or "Off-Site"
- is_key_position: true for managers/leads/key personnel, false otherwise
- ftes: Number of FTEs - EXTRACT from document if available, or calculate from hours ÷ 1920
- hours_per_year: Dict {"1": hours, "2": hours, ...} - EXTRACT EXACT numbers from document tables

IMPORTANT: If document has a labor hours table, use those EXACT numbers. Do not estimate.""",

        """DECISION FLOW:
1. Scan entire document for tables with hours/FTE data
2. IF tables found with hours → Extract exact hours for each position and year
3. IF tables found with FTEs → Extract FTEs, calculate hours as FTE × 1920
4. IF no numerical data found → THEN use web search to research typical staffing levels

WEB SEARCH TRIGGERS (only if no data in document):
- PWS/SOW documents with only job descriptions
- 24/7 operations without staffing numbers
- Roles without any quantitative data

Remember: Document data > Web research. Extract first, estimate only when necessary.""",
    ]

    agent = Agent(
        name="Contract Staffing Parser Agent",
        model=llm,
        tools=[exa_tool],
        instructions=instructions,
        markdown=True,
        debug_mode=True,
    )

    return agent


def _validate_result(
    result: Dict[str, Any],
    default_years: int,
    default_fte_hours: int
) -> Dict[str, Any]:
    """Validate and fix the result."""

    metadata = result.get('metadata', {})
    total_years = metadata.get('total_years') or default_years
    standard_fte_hours = metadata.get('standard_fte_hours') or default_fte_hours

    metadata['total_years'] = total_years
    metadata['standard_fte_hours'] = standard_fte_hours

    positions = result.get('positions', [])
    validated_positions = []

    for pos in positions:
        labor_category = pos.get('labor_category', 'Unknown Position')
        ftes = pos.get('ftes') or 1
        experience = pos.get('experience') or 3

        hours_per_year = pos.get('hours_per_year', {})
        if not hours_per_year:
            hours_per_year = {}

        for y in range(1, total_years + 1):
            year_key = str(y)
            if year_key not in hours_per_year or not hours_per_year[year_key]:
                hours_per_year[year_key] = ftes * standard_fte_hours
            else:
                hours_per_year[year_key] = int(hours_per_year[year_key])

        validated_positions.append({
            'labor_category': labor_category,
            'description': pos.get('description'),
            'experience': int(experience) if experience else 3,
            'location': pos.get('location') or metadata.get('location'),
            'location_type': pos.get('location_type') or 'On-Site',
            'is_key_position': bool(pos.get('is_key_position', False)),
            'ftes': int(ftes),
            'hours': int(ftes) * standard_fte_hours,
            'hours_per_year': hours_per_year,
        })

    travel = []
    for t in result.get('travel', []):
        if isinstance(t, dict) and t.get('amount_per_year'):
            travel.append({
                'description': t.get('description', 'Travel'),
                'amount_per_year': {str(k): float(v) for k, v in t['amount_per_year'].items()},
            })
        elif isinstance(t, str):
            # Just a description without amounts
            travel.append({
                'description': t,
                'amount_per_year': {},
            })

    odcs = []
    for o in result.get('odcs', []):
        if isinstance(o, dict) and o.get('amount_per_year'):
            odcs.append({
                'category': o.get('category', 'Other'),
                'description': o.get('description'),
                'amount_per_year': {str(k): float(v) for k, v in o['amount_per_year'].items()},
            })
        elif isinstance(o, str):
            odcs.append({
                'category': 'Other',
                'description': o,
                'amount_per_year': {},
            })

    return {
        'metadata': metadata,
        'positions': validated_positions,
        'travel': travel,
        'odcs': odcs,
    }


async def parse_pdf_smart(
    file_path: str,
    default_fte_hours: int = 1920,
    default_years: int = 5
) -> Dict[str, Any]:
    """
    Smart document parser using Agent with Exa web search.

    Args:
        file_path: Path to document (PDF, DOCX, TXT)
        default_fte_hours: Default FTE hours per year
        default_years: Default contract duration

    Returns:
        Dict with metadata, positions, travel, odcs
    """
    print(f"\n{'='*60}")
    print(f"Smart PDF Parser - Agent with Web Research")
    print(f"{'='*60}")
    print(f"\nFile: {Path(file_path).name}")

    # Step 1: Extract full text
    print(f"\n📄 Step 1: Extracting document text...")
    document_text = _extract_text(file_path)
    print(f"  ✓ Extracted {len(document_text):,} characters")

    # Step 2: Create agent and parse
    print(f"\n🤖 Step 2: Creating parsing agent with Exa web search...")
    agent = _create_parsing_agent()

    prompt = f"""Parse this government contract document and extract ALL staffing information.

DOCUMENT TEXT:
{document_text}

CRITICAL INSTRUCTIONS:
1. FIRST: Search the document for any tables with labor hours, FTEs, or level of effort data
2. If you find hours/FTE data in tables → Extract EXACT numbers as shown in document
3. ONLY use web search if the document has NO numerical staffing data at all
4. Extract ALL positions/labor categories mentioned

DEFAULT PARAMETERS (use only if not specified in document):
- Contract Duration: {default_years} years
- Standard FTE Hours: {default_fte_hours} hours/year

OUTPUT FORMAT - Return ONLY this JSON structure:
{{
  "metadata": {{
    "project_name": "...",
    "location": "State name",
    "base_years": 1,
    "option_years": 4,
    "total_years": 5,
    "standard_fte_hours": {default_fte_hours}
  }},
  "positions": [
    {{
      "labor_category": "Position Title",
      "description": "Job description...",
      "experience": 5,
      "location": "State name",
      "location_type": "On-Site",
      "is_key_position": false,
      "ftes": 4,
      "hours_per_year": {{"1": 7680, "2": 7680, "3": 7680, "4": 7680, "5": 7680}}
    }}
  ],
  "travel": [],
  "odcs": []
}}

IMPORTANT RULES:
- If document has hours table → Use EXACT hours from table
- If document has FTE counts → Use exact FTEs, calculate hours = FTEs × {default_fte_hours}
- If NO numerical data in document → Then use web search to research typical staffing
- hours_per_year keys should be "1", "2", "3", etc. for each contract year

Return ONLY the JSON, no other text."""

    print(f"\n🔍 Step 3: Running agent (may perform web searches)...")

    # Run the agent
    response = await agent.arun(prompt)

    # Extract JSON from response
    response_text = response.content if hasattr(response, 'content') else str(response)

    # Clean up response to get JSON
    if '```json' in response_text:
        response_text = response_text.split('```json')[1].split('```')[0]
    elif '```' in response_text:
        response_text = response_text.split('```')[1].split('```')[0]

    response_text = response_text.strip()

    try:
        raw_result = json.loads(response_text)
    except json.JSONDecodeError as e:
        print(f"  ❌ JSON parsing error: {e}")
        print(f"  Response: {response_text[:500]}...")
        # Return empty result
        raw_result = {
            'metadata': {'total_years': default_years, 'standard_fte_hours': default_fte_hours},
            'positions': [],
            'travel': [],
            'odcs': []
        }

    # Step 4: Validate
    print(f"\n✅ Step 4: Validating results...")
    result = _validate_result(raw_result, default_years, default_fte_hours)

    # Summary
    print(f"\n{'='*60}")
    print(f"✅ Parsing Complete")
    print(f"{'='*60}")

    metadata = result['metadata']
    positions = result['positions']

    print(f"\n  Project: {metadata.get('project_name', 'Unknown')}")
    print(f"  Location: {metadata.get('location', 'Unknown')}")
    print(f"  Duration: {metadata.get('total_years', default_years)} years")

    print(f"\n  Positions: {len(positions)}")
    total_ftes = sum(p.get('ftes', 1) for p in positions)
    print(f"  Total FTEs: {total_ftes}")

    total_hours_y1 = sum(p.get('hours_per_year', {}).get('1', 0) for p in positions)
    print(f"  Year 1 Hours: {total_hours_y1:,}")

    print(f"\n  Positions extracted:")
    for i, pos in enumerate(positions[:10]):
        ftes = pos.get('ftes', 1)
        exp = pos.get('experience', '?')
        key = "🔑" if pos.get('is_key_position') else "  "
        print(f"    {key} {i+1}. {pos['labor_category'][:45]}")
        print(f"         FTEs: {ftes}, Exp: {exp}yrs")

    if len(positions) > 10:
        print(f"    ... and {len(positions) - 10} more positions")

    return result


def test_smart_parser(file_path: str) -> Dict[str, Any]:
    """Test the smart parser on a file."""
    import asyncio
    return asyncio.run(parse_pdf_smart(file_path))


if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        result = test_smart_parser(sys.argv[1])
        print(f"\n\n=== FINAL JSON ===")
        print(json.dumps(result, indent=2, default=str))
    else:
        print("Usage: python smart_pdf_parser.py <path_to_document>")
