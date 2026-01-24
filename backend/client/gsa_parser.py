"""GSA Contract parsing using LLM (GPT-4o/Claude)."""

import re
from typing import List, Optional
from pydantic import BaseModel, Field
import json
from app.settings import settings
from client.jd_parser import _convert_excel_to_csv
from client.llm_client import get_chat_llm


# =====================================================================
# METADATA SCHEMA
# =====================================================================

class GSAContractMetadata(BaseModel):
    """
    Metadata extraction for GSA contract (contract number, dates, company name).
    Labor categories are extracted separately using dual-parser approach.
    """
    contract_number: Optional[str] = Field(
        None,
        description="GSA contract number (e.g., 'GS-35F-0119Y', '47QRCA25DS242', '47QTCA20D003R')"
    )
    contract_start_date: Optional[str] = Field(
        None,
        description="Contract start date or period of performance start"
    )
    contract_end_date: Optional[str] = Field(
        None,
        description="Contract end date or period of performance end"
    )
    company_name: Optional[str] = Field(
        None,
        description="Contractor/company name"
    )
    year_columns: Optional[List[str]] = Field(
        None,
        description="List of year column headers from the rate table (e.g., ['Year 6', 'Year 7', 'Year 8', 'Year 9', 'Year 10'] or ['Year 1', 'Year 2', 'Year 3', 'Year 4', 'Year 5'])"
    )


# =====================================================================
# HELPER FUNCTIONS
# =====================================================================

def _convert_rtf_to_txt(rtf_path: str) -> str:
    """
    Convert RTF file to TXT for text extraction.

    Args:
        rtf_path: Path to the RTF file

    Returns:
        Path to temporary TXT file
    """
    import subprocess
    import tempfile

    temp_txt = tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False)
    temp_txt.close()

    try:
        # Use textutil (macOS) to convert RTF to TXT
        subprocess.run(
            ['textutil', '-convert', 'txt', '-output', temp_txt.name, rtf_path],
            check=True,
            capture_output=True
        )
        print(f"  Converted RTF to TXT: {temp_txt.name}")
        return temp_txt.name
    except (subprocess.CalledProcessError, FileNotFoundError):
        # Fallback: try striprtf library
        try:
            from striprtf.striprtf import rtf_to_text
            with open(rtf_path, 'r', encoding='utf-8', errors='ignore') as f:
                rtf_content = f.read()
            txt_content = rtf_to_text(rtf_content)
            with open(temp_txt.name, 'w', encoding='utf-8') as f:
                f.write(txt_content)
            print(f"  Converted RTF to TXT (striprtf): {temp_txt.name}")
            return temp_txt.name
        except ImportError:
            raise ValueError("Cannot convert RTF. Install striprtf: pip install striprtf")


def _extract_full_text(file_path: str) -> str:
    """
    Extract full text from file (no chunking).

    Args:
        file_path: Path to file

    Returns:
        Full document text as single string
    """
    import os
    file_ext = os.path.splitext(file_path)[1].lower()

    # =========================================================================
    # PDF: Extract all pages
    # =========================================================================
    if file_ext == '.pdf':
        try:
            import PyPDF2
            pages = []
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        pages.append(text)

            full_text = '\n\n'.join(pages)
            print(f"     [PDF] Extracted {len(full_text):,} characters from {len(pages)} pages")
            return full_text
        except ImportError:
            raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")

    # =========================================================================
    # DOCX: Extract paragraphs and tables
    # =========================================================================
    elif file_ext == '.docx':
        try:
            import docx
            doc = docx.Document(file_path)
            all_content = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    all_content.append(para.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        all_content.append(row_text)

            full_text = '\n'.join(all_content)
            print(f"     [DOCX] Extracted {len(full_text):,} characters")
            return full_text
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

    # =========================================================================
    # CSV: Extract all rows
    # =========================================================================
    elif file_ext == '.csv':
        import csv
        all_lines = []

        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            reader = csv.reader(f)
            for row in reader:
                line = ' | '.join([cell.strip() for cell in row if cell.strip()])
                if line:
                    all_lines.append(line)

        full_text = '\n'.join(all_lines)
        print(f"     [CSV] Extracted {len(full_text):,} characters from {len(all_lines)} rows")
        return full_text

    # =========================================================================
    # TXT/RTF: Extract all text
    # =========================================================================
    else:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            full_text = f.read()

        print(f"     [TXT] Extracted {len(full_text):,} characters")
        return full_text




def _extract_metadata_with_llm(full_text: str) -> GSAContractMetadata:
    """
    Extract metadata from full document text using LLM.

    Args:
        full_text: Complete document text

    Returns:
        GSAContractMetadata instance
    """
    llm = get_chat_llm(model="gpt-4.1",api_key=settings.OPENAI_API_KEY,base_url="https://api.openai.com/v1", max_tokens=10000)

    prompt = f"""Extract metadata from this GSA contract document.

FIELDS TO EXTRACT:
1. contract_number: GSA contract number (e.g., 'GS-35F-0119Y', '47QRCA25DS242', '47QTCA20D003R')
2. contract_start_date: Contract start date or period of performance start
3. contract_end_date: Contract end date or period of performance end
4. company_name: Contractor/company name
5. year_columns: List of year column headers from rate tables (e.g., ['Year 6', 'Year 7', 'Year 8'] or ['Year 1', 'Year 2'])

Look at the BEGINNING of the document for contract number, dates, and company name.
Look at RATE TABLES for year column headers.

Return ONLY valid JSON in this format:
{{
  "contract_number": "GS-35F-0119Y",
  "contract_start_date": "January 1, 2025",
  "contract_end_date": "December 31, 2030",
  "company_name": "ACME Corporation",
  "year_columns": ["Year 6", "Year 7", "Year 8", "Year 9", "Year 10"]
}}

If a field is not found, use null.

DOCUMENT TEXT :
{full_text}

Return ONLY valid JSON:"""

    max_retries = 3
    for retry_count in range(max_retries):
        try:
            # Add retry guidance to prompt if this is a retry
            current_prompt = prompt
            if retry_count > 0:
                current_prompt = f"{prompt}\n\nIMPORTANT: Your previous response was malformed. Please provide COMPLETE valid JSON with all brackets and quotes closed properly."
                print(f"     [Metadata] 🔄 Retry {retry_count}/{max_retries - 1} due to JSON error...")

            response = llm.invoke(current_prompt)
            response_text = response.content.strip()

            # Remove markdown code blocks if present
            if response_text.startswith('```'):
                # Find the first ``` and last ```
                parts = response_text.split('```')
                if len(parts) >= 3:
                    response_text = parts[1]
                    if response_text.startswith('json'):
                        response_text = response_text[4:]
                response_text = response_text.strip()

            # Additional cleanup for common JSON issues
            # Remove any trailing commas before ] or }
            response_text = re.sub(r',\s*([}\]])', r'\1', response_text)

            # Try to find JSON object in response if it's wrapped in other text
            if not response_text.startswith('{'):
                match = re.search(r'\{.*\}', response_text, re.DOTALL)
                if match:
                    response_text = match.group(0)

            # Parse JSON
            metadata_dict = json.loads(response_text)
            metadata = GSAContractMetadata(**metadata_dict)

            print(f"     [Metadata] ✓ Extracted via LLM")
            if metadata.year_columns:
                print(f"     [Metadata] Year columns: {metadata.year_columns}")
            else:
                print(f"     [Metadata] ⚠️  No year columns detected")

            return metadata

        except json.JSONDecodeError as e:
            print(f"     [Metadata] ❌ JSON Error (attempt {retry_count + 1}/{max_retries}): {e}")
            print(f"     [Metadata] Response preview: {response_text[:500] if 'response_text' in locals() else 'N/A'}...")
            if retry_count >= max_retries - 1:
                print(f"     [Metadata] ❌ Failed after {max_retries} attempts")
                return GSAContractMetadata()
        except Exception as e:
            print(f"     [Metadata] ❌ Error: {e}")
            return GSAContractMetadata()

    return GSAContractMetadata()


def _extract_descriptions_with_llm(full_text: str) -> List[dict]:
    """
    Extract ONLY job descriptions, qualifications, and experience (NO RATES).

    Args:
        full_text: Complete document text

    Returns:
        List of dicts with title, sin, description, experience
    """
    llm = get_chat_llm(model="gpt-5-mini-2025-08-07",api_key=settings.OPENAI_API_KEY,base_url="https://api.openai.com/v1", max_tokens=30000)

    prompt = f"""Extract job descriptions and qualifications from this GSA contract document.

EXTRACTION RULES:
1. Extract ALL job titles you find
2. Extract descriptions and experience/education requirements
3. IGNORE dollar amounts and rates - we only want qualifications

FIELDS TO EXTRACT:
- title: Job title/position name (REQUIRED)
- sin: SIN code if present (e.g., "541330ENG", "541611")
- description: Job description, duties, responsibilities. Look for:
  * Text paragraphs describing what the position does
  * Duty statements or role descriptions
  * Required skills or competencies
  * If NO description exists, omit this field
- experience: Years of experience or education required. Look for:
  * Text like "5 years", "3-5 years minimum", "10+ years"
  * Education: "Bachelor's degree", "Master's required", "PhD preferred"
  * Level indicators: "Senior" (5+ years), "Junior" (1-3 years), "Lead" (7+ years)
  * Roman numerals: "I" = Entry, "II" = 2-4 years, "III" = 5-7 years, "IV" = 8+ years
  * If NO experience info exists, omit this field

OUTPUT FORMAT:
[
  {{
    "title": "Senior Systems Engineer",
    "sin": "541330",
    "description": "Designs and implements enterprise infrastructure solutions...",
    "experience": "5+ years with Bachelor's degree"
  }},
  {{
    "title": "Administrative Specialist",
    "sin": "541611",
    "description": "Provides administrative support including scheduling..."
  }}
]

DOCUMENT TEXT:
{full_text}

Return ONLY valid JSON array. If no job descriptions found, return empty array []:"""

    max_retries = 3
    for retry_count in range(max_retries):
        try:
            # Add retry guidance to prompt if this is a retry
            current_prompt = prompt
            if retry_count > 0:
                current_prompt = f"{prompt}\n\nIMPORTANT: Your previous response was incomplete or malformed. Please provide the COMPLETE valid JSON array with all entries fully formed. Do not truncate the output."
                print(f"     [Descriptions] 🔄 Retry {retry_count}/{max_retries - 1} due to JSON error...")

            response = llm.invoke(current_prompt)
            response_text = response.content.strip()

            # Remove markdown code blocks if present
            if response_text.startswith('```'):
                # Find the first ``` and last ```
                parts = response_text.split('```')
                if len(parts) >= 3:
                    response_text = parts[1]
                    if response_text.startswith('json'):
                        response_text = response_text[4:]
                response_text = response_text.strip()

            # Additional cleanup for common JSON issues
            # Remove any trailing commas before ] or }
            response_text = re.sub(r',\s*([}\]])', r'\1', response_text)

            # Try to find JSON array in response if it's wrapped in other text
            if not response_text.startswith('['):
                match = re.search(r'\[.*\]', response_text, re.DOTALL)
                if match:
                    response_text = match.group(0)

            # Parse JSON
            descriptions = json.loads(response_text)

            if not isinstance(descriptions, list):
                descriptions = [descriptions]

            print(f"     [Descriptions] ✓ Extracted {len(descriptions)} job descriptions")
            return descriptions

        except json.JSONDecodeError as e:
            print(f"     [Descriptions] ❌ JSON Error (attempt {retry_count + 1}/{max_retries}): {e}")
            print(f"     [Descriptions] Response preview: {response_text[:500] if 'response_text' in locals() else 'N/A'}...")
            if retry_count >= max_retries - 1:
                print(f"     [Descriptions] ❌ Failed after {max_retries} attempts")
                return []
        except Exception as e:
            print(f"     [Descriptions] ❌ Error: {e}")
            return []

    return []


def _extract_rates_with_llm(full_text: str, year_columns: Optional[List[str]] = None) -> List[dict]:
    """
    Extract ONLY rate tables with dollar amounts (NO DESCRIPTIONS).

    Args:
        full_text: Complete document text
        year_columns: List of year column headers from metadata

    Returns:
        List of dicts with title, sin, rates_by_year
    """
    llm = get_chat_llm(model="gpt-5-mini-2025-08-07",api_key=settings.OPENAI_API_KEY,base_url="https://api.openai.com/v1", max_tokens=30000)

    # Build year context from metadata
    if year_columns and len(year_columns) > 0:
        year_nums = []
        for col in year_columns:
            match = re.search(r'\d+', col)
            if match:
                year_nums.append(match.group())

        year_context = f"""
🎯 IMPORTANT: This document has year columns: {', '.join(year_columns)}
Use these EXACT year numbers as keys in rates_by_year: {', '.join(year_nums)}
"""
    else:
        year_context = """
Look at the table header row to find year columns (e.g., "Year 1", "Year 6", etc.)
Use the EXACT year numbers from headers as keys in rates_by_year.
"""

    prompt = f"""Extract ONLY rate/pricing tables from this GSA contract document.
{year_context}

🎯 YOUR TASK: Find tables with job titles and dollar amounts (hourly rates).

EXTRACTION RULES:
1. Look for TABLE FORMAT with:
   - Column 1: Job titles (e.g., "Senior Systems Engineer", "Administrative Specialist")
   - Other columns: Dollar amounts per year (e.g., "$125.50", "$45.00")
2. ONLY extract rows that have BOTH:
   - A clear job title/labor category name
   - At least one dollar rate (hourly or annual)
3. IGNORE completely:
   - Job descriptions and duty statements
   - Qualifications, experience, education requirements
   - Table headers and SIN category descriptions
   - Rows without any dollar amounts

DOLLAR AMOUNT HANDLING:
- If you see hourly rates (e.g., "$125.50/hr", "$45.00"), use them directly
- If you see annual salaries (e.g., "$95,000"), convert to hourly: annual ÷ 2080
- Remove any "$" signs and commas from the numbers
- Store rates as plain numbers (e.g., 125.50, not "$125.50")

YEAR COLUMN DETECTION:
- Look at the table header row to identify year columns
- Common formats: "Year 1", "Year 6", "Yr 1", "Option Year 2", "Base Year"
- Extract the year NUMBER from the header (e.g., "Year 6" → use key "6")
- If no year numbers visible, use sequential numbers: "1", "2", "3"...

FIELDS TO EXTRACT:
- title: Exact job title from the table (REQUIRED)
- sin: SIN code if visible in same row (e.g., "541330ENG", "541611") - optional
- rates_by_year: Dictionary with year numbers as STRING keys and rates as NUMBERS
  Example: {{"1": 125.50, "2": 128.00, "3": 130.56}}

OUTPUT FORMAT EXAMPLES:

Example 1 - Hourly rates:
[
  {{
    "title": "Senior Systems Engineer",
    "sin": "541330",
    "rates_by_year": {{"1": 125.50, "2": 128.00, "3": 130.56}}
  }}
]

Example 2 - Annual salary converted:
[
  {{
    "title": "Administrative Specialist",
    "rates_by_year": {{"1": 45.67, "2": 46.35}}
  }}
]

Example 3 - Multiple years:
[
  {{
    "title": "Program Manager",
    "sin": "541611",
    "rates_by_year": {{"6": 175.00, "7": 180.25, "8": 185.65, "9": 191.22, "10": 197.00}}
  }}
]

IMPORTANT:
- Focus ONLY on finding rate tables (rows with $ amounts)
- Ignore any text paragraphs, descriptions, or qualifications
- If this document has no rate tables, return empty array: []

DOCUMENT TEXT:
{full_text}

Return ONLY valid JSON array:"""

    max_retries = 3
    for retry_count in range(max_retries):
        try:
            # Add retry guidance to prompt if this is a retry
            current_prompt = prompt
            if retry_count > 0:
                current_prompt = f"{prompt}\n\nIMPORTANT: Your previous response was truncated or malformed. Please provide the COMPLETE valid JSON array. Ensure ALL labor categories are included and the JSON is not cut off."
                print(f"     [Rates] 🔄 Retry {retry_count}/{max_retries - 1} due to JSON error...")

            response = llm.invoke(current_prompt)
            response_text = response.content.strip()

            # Remove markdown code blocks if present
            if response_text.startswith('```'):
                # Find the first ``` and last ```
                parts = response_text.split('```')
                if len(parts) >= 3:
                    response_text = parts[1]
                    if response_text.startswith('json'):
                        response_text = response_text[4:]
                response_text = response_text.strip()

            # Additional cleanup for common JSON issues
            # Remove any trailing commas before ] or }
            response_text = re.sub(r',\s*([}\]])', r'\1', response_text)

            # Try to find JSON array in response if it's wrapped in other text
            if not response_text.startswith('['):
                match = re.search(r'\[.*\]', response_text, re.DOTALL)
                if match:
                    response_text = match.group(0)

            # Parse JSON
            rates = json.loads(response_text)

            if not isinstance(rates, list):
                rates = [rates]

            print(f"     [Rates] ✓ Extracted {len(rates)} rate entries")
            return rates

        except json.JSONDecodeError as e:
            print(f"     [Rates] ❌ JSON Error (attempt {retry_count + 1}/{max_retries}): {e}")
            print(f"     [Rates] Response preview: {response_text[:500] if 'response_text' in locals() else 'N/A'}...")
            if retry_count >= max_retries - 1:
                print(f"     [Rates] ❌ Failed after {max_retries} attempts")
                return []
        except Exception as e:
            print(f"     [Rates] ❌ Error: {e}")
            return []

    return []


def _merge_descriptions_and_rates(descriptions: List[dict], rates: List[dict]) -> List[dict]:
    """
    Merge job descriptions with rate tables by matching title + SIN.

    Args:
        descriptions: List of dicts with {title, sin, description, experience}
        rates: List of dicts with {title, sin, rates_by_year}

    Returns:
        Merged list with complete labor category data
    """
    from difflib import SequenceMatcher

    def normalize_title(title: str) -> str:
        """Normalize title for fuzzy matching."""
        return title.lower().strip().replace('-', ' ').replace('/', ' ')

    def fuzzy_match_score(str1: str, str2: str) -> float:
        """Calculate similarity score between two strings (0.0 to 1.0)."""
        return SequenceMatcher(None, normalize_title(str1), normalize_title(str2)).ratio()

    # Build lookup for rates by (title, sin) and by title only
    rates_by_key = {}
    rates_by_title = {}

    for rate in rates:
        title = rate.get('title', '').strip()
        sin = rate.get('sin', '').strip()

        if not title:
            continue

        # Index by (title, sin) if SIN exists
        if sin:
            key = (normalize_title(title), sin.upper())
            rates_by_key[key] = rate

        # Index by title only
        rates_by_title[normalize_title(title)] = rate

    # Merge descriptions with rates
    merged = []
    matched_rate_indices = set()  # Track indices of matched rates

    for desc in descriptions:
        title = desc.get('title', '').strip()
        sin = desc.get('sin', '').strip()

        if not title:
            continue

        # Try exact match by (title, sin)
        matched_rate = None
        if sin:
            key = (normalize_title(title), sin.upper())
            matched_rate = rates_by_key.get(key)

        # Try exact match by title only
        if not matched_rate:
            matched_rate = rates_by_title.get(normalize_title(title))

        # Try fuzzy match if no exact match
        matched_rate_index = None
        if not matched_rate:
            best_score = 0.0
            best_match = None
            best_match_index = None

            for idx, rate in enumerate(rates):
                rate_title = rate.get('title', '').strip()
                if not rate_title:
                    continue

                score = fuzzy_match_score(title, rate_title)
                if score > best_score and score >= 0.85:  # 85% similarity threshold
                    best_score = score
                    best_match = rate
                    best_match_index = idx

            if best_match:
                matched_rate = best_match
                matched_rate_index = best_match_index
                print(f"     [Merge] Fuzzy matched: '{title}' → '{best_match.get('title')}' (score: {best_score:.2f})")
        else:
            # Find index of matched rate
            for idx, rate in enumerate(rates):
                if rate == matched_rate:
                    matched_rate_index = idx
                    break

        # Merge data
        if matched_rate:
            merged_entry = {
                'title': title,  # Use description's title (usually cleaner)
                'sin': sin or matched_rate.get('sin'),
                'description': desc.get('description'),
                'experience': desc.get('experience'),
                'rates_by_year': matched_rate.get('rates_by_year', {})
            }
            merged.append(merged_entry)
            if matched_rate_index is not None:
                matched_rate_indices.add(matched_rate_index)
        else:
            # Keep description even without rates
            merged.append(desc)
            print(f"     [Merge] ⚠️  No rates found for: '{title}'")

    # Add unmatched rates (rates without descriptions)
    for idx, rate in enumerate(rates):
        if idx not in matched_rate_indices:
            title = rate.get('title', '').strip()
            if title:
                merged.append({
                    'title': title,
                    'sin': rate.get('sin'),
                    'rates_by_year': rate.get('rates_by_year', {})
                })
                print(f"     [Merge] ⚠️  No description found for: '{title}'")

    return merged


# =====================================================================
# MAIN FUNCTION
# =====================================================================

def parse_gsa_contract(file_path: str) -> dict:
    """
    Parse GSA contract and return data for MongoDB storage.

    Args:
        file_path: Path to GSA contract (PDF, Excel, or RTF)

    Returns:
        Dict with contract_number, dates, labor_categories, needs_date
    """
    import os

    file_ext = os.path.splitext(file_path)[1].lower()
    temp_file_path = None

    # Handle Excel files
    if file_ext in ['.xlsx', '.xls']:
        print(f"  Converting Excel to CSV...")
        temp_file_path = _convert_excel_to_csv(file_path)
        file_path = temp_file_path

    # Handle RTF files
    elif file_ext == '.rtf':
        print(f"  Converting RTF to TXT...")
        temp_file_path = _convert_rtf_to_txt(file_path)
        file_path = temp_file_path

    try:
        # ========================================================================
        # STEP 1: Extract full document text (no chunking)
        # ========================================================================
        print("  📄 Extracting full document...")
        full_text = _extract_full_text(file_path)

        if not full_text.strip():
            raise ValueError("Failed to extract text from document")

        # ========================================================================
        # STEP 2: Extract metadata first (needed for year_columns)
        # ========================================================================
        print("  🔍 Step 1: Extracting metadata...")
        metadata = _extract_metadata_with_llm(full_text)
        year_columns = metadata.year_columns or []

        # ========================================================================
        # STEP 3: Parallel extraction of descriptions + rates (using metadata)
        # ========================================================================
        import concurrent.futures as cf

        print("  🚀 Step 2: Dual-parser extraction (2 parallel workers)...")
        print("     Worker 1: Job descriptions")
        print("     Worker 2: Rate tables")

        # Run descriptions and rates extraction in parallel
        with cf.ThreadPoolExecutor(max_workers=2) as executor:
            desc_future = executor.submit(_extract_descriptions_with_llm, full_text)
            rates_future = executor.submit(_extract_rates_with_llm, full_text, year_columns)

            all_descriptions = desc_future.result()
            all_rates = rates_future.result()

        print(f"     [Descriptions] Found {len(all_descriptions)} entries")
        print(f"     [Rates] Found {len(all_rates)} entries")

        # ========================================================================
        # STEP 4: MERGE descriptions + rates by title/SIN matching
        # ========================================================================
        print(f"  🔗 Step 3: Merging descriptions with rates...")
        merged = _merge_descriptions_and_rates(all_descriptions, all_rates)
        print(f"     ✓ Merged: {len(merged)} labor categories")

        # ========================================================================
        # STEP 5: Deduplicate and format for storage
        # ========================================================================
        seen_titles = set()
        labor_categories = []

        for cat in merged:
            title = cat.get('title', '').strip()
            title_lower = title.lower()

            if title and title_lower not in seen_titles:
                seen_titles.add(title_lower)

                labor_categories.append({
                    "lcat_id": f"lcat_{len(labor_categories)}",
                    "sin": cat.get('sin'),
                    "title": title,
                    "description": cat.get('description'),
                    "experience": cat.get('experience'),
                    "rates_by_year": cat.get('rates_by_year', {})
                })

        # ========================================================================
        # RESULTS SUMMARY
        # ========================================================================
        print(f"\n  ✅ Extraction Complete:")
        print(f"     Contract Number: {metadata.contract_number}")
        print(f"     Company Name: {metadata.company_name}")
        print(f"     Start Date: {metadata.contract_start_date}")
        print(f"     End Date: {metadata.contract_end_date}")
        print(f"     Total Labor Categories: {len(labor_categories)}")

        # Show first 3 and last 3
        if labor_categories:
            print(f"\n     Sample categories:")
            for i, lcat in enumerate(labor_categories[:3]):
                rates = lcat.get('rates_by_year', {})
                rate_str = f"${list(rates.values())[0]:.2f}" if rates else "No rates"
                print(f"       [{i+1}] {lcat['title']} - {rate_str}")

            if len(labor_categories) > 6:
                print(f"       ...")
                for i, lcat in enumerate(labor_categories[-3:]):
                    idx = len(labor_categories) - 3 + i + 1
                    rates = lcat.get('rates_by_year', {})
                    rate_str = f"${list(rates.values())[0]:.2f}" if rates else "No rates"
                    print(f"       [{idx}] {lcat['title']} - {rate_str}")

        # Parse dates
        start_date = _parse_date(metadata.contract_start_date)
        end_date = _parse_date(metadata.contract_end_date)

        return {
            "contract_number": metadata.contract_number,
            "contract_start_date": start_date,
            "contract_end_date": end_date,
            "company_name": metadata.company_name,
            "labor_categories": labor_categories,
            "needs_date": start_date is None
        }

    finally:
        if temp_file_path:
            try:
                os.unlink(temp_file_path)
            except Exception:
                pass


def _parse_date(date_str: Optional[str]) -> Optional[str]:
    """Parse date string to ISO format."""
    if not date_str:
        return None
    try:
        from dateutil import parser
        return parser.parse(date_str).strftime("%Y-%m-%d")
    except Exception:
        return date_str
